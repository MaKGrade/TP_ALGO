# Import des bibliothèques nécessaires
from abc import ABC, abstractmethod
from typing import Dict, List, Set, Tuple
import os
import re
import hashlib
from collections import defaultdict
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
import numpy as np
from tqdm import tqdm
import logging
from bs4 import BeautifulSoup
import markdown
import docx
import PyPDF2
import esprima
from ast import parse
import math
from functools import lru_cache
import pandas as pd
import ast
import pdfplumber

# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Classes de base pour les extracteurs
class BaseExtractor(ABC):
    @abstractmethod
    def can_handle(self, file_path: str) -> bool:
        pass
        
    @abstractmethod
    def extract(self, file_path: str) -> str:
        pass
        
    @abstractmethod
    def get_metadata(self, file_path: str) -> Dict:
        pass

class TextExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith('.txt')
    
    def extract(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                logger.info(f"Contenu extrait de {file_path}: {content[:100]}...")  # Log les 100 premiers caractères
                return content
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
                logger.info(f"Contenu extrait (Latin-1) de {file_path}: {content[:100]}...")
                return content
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction de {file_path}: {str(e)}")
            return ""
    
    def get_metadata(self, file_path: str) -> Dict:
        return {
            'type': 'text',
            'size': os.path.getsize(file_path)
        }

class DocExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith(('.doc', '.docx'))
    
    def extract(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return '\n'.join(text)
    
    def get_metadata(self, file_path: str) -> Dict:
        doc = docx.Document(file_path)
        return {
            'type': 'document',
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'size': os.path.getsize(file_path)
        }

class PDFExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith('.pdf')
    
    def extract(self, file_path: str) -> str:
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text.append(extracted_text)
        return '\n'.join(text)
    
    def get_metadata(self, file_path: str) -> Dict:
        with pdfplumber.open(file_path) as pdf:
            return {
                'type': 'pdf',
                'pages': len(pdf.pages),
                'size': os.path.getsize(file_path)
            }

class LatexExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith('.tex')
    
    def extract(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            tex_content = f.read()
        return latex2text.latex2text(tex_content)
    
    def get_metadata(self, file_path: str) -> Dict:
        return {
            'type': 'latex',
            'size': os.path.getsize(file_path)
        }

class HTMLMarkdownExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith(('.html', '.md'))
    
    def extract(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        if file_path.endswith('.md'):
            html_content = markdown.markdown(content)
        else:
            html_content = content
            
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text(separator='\n')
    
    def get_metadata(self, file_path: str) -> Dict:
        return {
            'type': 'markup',
            'format': 'markdown' if file_path.endswith('.md') else 'html',
            'size': os.path.getsize(file_path)
        }

class ExcelExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith(('.xls', '.xlsx'))
    
    def extract(self, file_path: str) -> str:
        df = pd.read_excel(file_path, sheet_name=None)
        text = []
        for sheet_name, sheet_data in df.items():
            text.append(f"Sheet: {sheet_name}")
            text.append(sheet_data.to_string())
        return '\n\n'.join(text)
    
    def get_metadata(self, file_path: str) -> Dict:
        df = pd.read_excel(file_path, sheet_name=None)
        return {
            'type': 'spreadsheet',
            'sheets': len(df),
            'size': os.path.getsize(file_path)
        }

class PowerPointExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith(('.ppt', '.pptx'))
    
    def extract(self, file_path: str) -> str:
        prs = Presentation(file_path)
        text = []
        for slide in prs.slides:
            slide_text = []
            for shape in slide.shapes:
                if hasattr(shape, "text"):
                    slide_text.append(shape.text)
            text.append('\n'.join(slide_text))
        return '\n\n'.join(text)
    
    def get_metadata(self, file_path: str) -> Dict:
        prs = Presentation(file_path)
        return {
            'type': 'presentation',
            'slides': len(prs.slides),
            'size': os.path.getsize(file_path)
        }

class CodeExtractor(BaseExtractor):
    def can_handle(self, file_path: str) -> bool:
        return file_path.endswith(('.py', '.java', '.cpp', '.js', '.ts'))
    
    def extract(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Supprimer les commentaires selon le type de fichier
        ext = os.path.splitext(file_path)[1]
        if ext == '.py':
            tree = ast.parse(content)
            return self._extract_python_content(tree)
        elif ext == '.java':
            tree = javalang.parse.parse(content)
            return self._extract_java_content(tree)
        elif ext in ['.js', '.ts']:
            tree = esprima.parseScript(content)
            return self._extract_js_content(tree)
        elif ext == '.cpp':
            return self._clean_cpp_comments(content)
        
        return content
    
    def _extract_python_content(self, tree) -> str:
        """Extrait le contenu Python en ignorant les docstrings et commentaires"""
        code_elements = []
        for node in ast.walk(tree):
            if isinstance(node, (ast.FunctionDef, ast.ClassDef)):
                code_elements.append(node.name)
            elif isinstance(node, ast.Name):
                code_elements.append(node.id)
        return ' '.join(code_elements)
    
    def _extract_java_content(self, tree) -> str:
        """Extrait le contenu Java en ignorant les commentaires"""
        code_elements = []
        for path, node in tree:
            if isinstance(node, (javalang.tree.MethodDeclaration, javalang.tree.ClassDeclaration)):
                code_elements.append(node.name)
        return ' '.join(code_elements)
    
    def _extract_js_content(self, tree) -> str:
        """Extrait le contenu JavaScript/TypeScript en ignorant les commentaires"""
        code_elements = []
        def visit(node):
            if 'type' in node:
                if node['type'] in ['FunctionDeclaration', 'ClassDeclaration']:
                    code_elements.append(node.get('id', {}).get('name', ''))
                for key, value in node.items():
                    if isinstance(value, dict):
                        visit(value)
                    elif isinstance(value, list):
                        for item in value:
                            if isinstance(item, dict):
                                visit(item)
        visit(tree.toDict())
        return ' '.join(code_elements)
    
    def _clean_cpp_comments(self, content: str) -> str:
        """Nettoie les commentaires C++"""
        # Supprime les commentaires sur une ligne
        content = re.sub(r'//.*?\n', '\n', content)
        # Supprime les commentaires multi-lignes
        content = re.sub(r'/\*.*?\*/', '', content, flags=re.DOTALL)
        return content
    
    def get_metadata(self, file_path: str) -> Dict:
        ext = os.path.splitext(file_path)[1]
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            'type': 'code',
            'language': ext[1:],  # Retire le point
            'lines': len(content.splitlines()),
            'size': os.path.getsize(file_path)
        }

# Registre des extracteurs
class ExtractorRegistry:
    def __init__(self):
        self._extractors = []
        self._initialize_registry()
    
    def _initialize_registry(self):
        extractors = [
            TextExtractor(),
            DocExtractor(),
            PDFExtractor(),
            LatexExtractor(),
            HTMLMarkdownExtractor(),
            ExcelExtractor(),
            PowerPointExtractor(),
            CodeExtractor()
        ]
        for extractor in extractors:
            self.register(extractor)
    
    def register(self, extractor: BaseExtractor):
        self._extractors.append(extractor)
    
    def get_extractor(self, file_path: str) -> BaseExtractor:
        for extractor in self._extractors:
            if extractor.can_handle(file_path):
                return extractor
        raise ValueError(f"No extractor found for {file_path}")

class DocumentProcessor:
    def __init__(self):
        self.registry = ExtractorRegistry()
    
    def clean_text(self, text: str) -> str:
        text = text.lower()
        text = re.sub(r'[^\w\s]', '', text)
        text = ' '.join(text.split())
        return text
    
    def tokenize(self, text: str) -> List[List[str]]:
        sentences = sent_tokenize(text)
        return [word_tokenize(sent) for sent in sentences]

    def process_document(self, file_path: str) -> Tuple[str, List[List[str]]]:
        logger.info(f"Début du traitement du document : {file_path}")
        try:
            extractor = self.registry.get_extractor(file_path)
            logger.info(f"Utilisation de l'extracteur : {type(extractor).__name__}")
            
            text = extractor.extract(file_path)
            logger.info(f"Taille du texte extrait : {len(text)} caractères")
            
            if not text.strip():
                logger.warning(f"Le fichier {file_path} est vide ou n'a pas pu être extrait.")
                return "", []
                
            cleaned_text = self.clean_text(text)
            tokens = self.tokenize(cleaned_text)
            
            return cleaned_text, tokens
        except Exception as e:
            logger.error(f"Erreur lors du traitement du document {file_path}: {str(e)}")
            raise
# Algorithmes de détection
class SimilarityEngine:
    def __init__(self, window_size=3, ngram_size=2):
        self.window_size = window_size
        self.ngram_size = ngram_size
    
    @lru_cache(maxsize=1024)
    def generate_fingerprint(self, text: str) -> Set[str]:
        fingerprints = set()
        words = text.split()
        for i in range(len(words) - self.window_size + 1):
            window = ' '.join(words[i:i+self.window_size])
            hash_value = hashlib.md5(window.encode()).hexdigest()
            fingerprints.add(hash_value)
        return fingerprints
    
    def generate_ngrams(self, text: str) -> List[str]:
        words = text.split()
        ngrams = []
        for i in range(len(words) - self.ngram_size + 1):
            ngram = ' '.join(words[i:i+self.ngram_size])
            ngrams.append(ngram)
        return ngrams

    def _calculate_fingerprint_score(self, doc1: str, doc2: str) -> float:
        """Calcule le score de similarité basé sur les empreintes"""
        set1 = self.generate_fingerprint(doc1)
        set2 = self.generate_fingerprint(doc2)
        if not set1 or not set2:
            return 0.0
        intersection = set1.intersection(set2)
        return len(intersection) / max(len(set1), len(set2))

    def _calculate_ngram_score(self, doc1: str, doc2: str) -> float:
        """Calcule le score de similarité basé sur les n-grammes"""
        ngrams1 = set(self.generate_ngrams(doc1))
        ngrams2 = set(self.generate_ngrams(doc2))
        if not ngrams1 or not ngrams2:
            return 0.0
        intersection = ngrams1.intersection(ngrams2)
        return len(intersection) / max(len(ngrams1), len(ngrams2))

    def _calculate_lcs_score(self, doc1: str, doc2: str) -> float:
        """Calcule le score de similarité basé sur la plus longue sous-séquence commune"""
        words1 = doc1.split()
        words2 = doc2.split()
        if not words1 or not words2:
            return 0.0
        lcs = self.lcs_length(doc1, doc2)
        return lcs / max(len(words1), len(words2))

    def lcs_length(self, text1: str, text2: str) -> int:
        """Calcule la longueur de la plus longue sous-séquence commune"""
        words1 = text1.split()
        words2 = text2.split()
        m, n = len(words1), len(words2)
        dp = [[0] * (n + 1) for _ in range(m + 1)]
        
        for i in range(1, m + 1):
            for j in range(1, n + 1):
                if words1[i-1] == words2[j-1]:
                    dp[i][j] = dp[i-1][j-1] + 1
                else:
                    dp[i][j] = max(dp[i-1][j], dp[i][j-1])
        
        return dp[m][n]

    def _get_weights(self, ext1: str, ext2: str) -> Dict:
        """Retourne les poids pour chaque métrique selon les types de fichiers"""
        if ext1.endswith(('.py', '.java', '.cpp', '.js', '.ts')) and ext2.endswith(('.py', '.java', '.cpp', '.js', '.ts')):
            return {
                'fingerprint': 0.3,
                'ngram': 0.3,
                'lcs': 0.2,
                'cosine': 0.2
            }
        else:
            return {
                'fingerprint': 0.25,
                'ngram': 0.25,
                'lcs': 0.25,
                'cosine': 0.25
            }

    def _get_threshold(self, ext1: str, ext2: str) -> float:
        """Détermine le seuil de plagiat selon les types de fichiers"""
        if ext1.endswith(('.py', '.java', '.cpp', '.js', '.ts')):
            return 0.6  # Seuil plus élevé pour le code
        elif ext1.endswith(('.html', '.md')):
            return 0.5  # Seuil moyen pour le markup
        else:
            return 0.5  # Seuil standard pour le texte

    def cosine_similarity(self, doc1: str, doc2: str) -> float:
        """Calcule la similarité cosinus entre deux documents"""
        words1 = doc1.split()
        words2 = doc2.split()
        
        vocabulary = list(set(words1 + words2))
        
        vector1 = [words1.count(word) for word in vocabulary]
        vector2 = [words2.count(word) for word in vocabulary]
        
        dot_product = sum(a * b for a, b in zip(vector1, vector2))
        norm1 = math.sqrt(sum(a * a for a in vector1))
        norm2 = math.sqrt(sum(b * b for b in vector2))
        
        if norm1 == 0 or norm2 == 0:
            return 0.0
            
        return dot_product / (norm1 * norm2)

    def calculate_similarity_score(self, doc1: str, doc2: str, ext1: str, ext2: str) -> Dict:
        """Calcule le score de similarité global entre deux documents"""
        try:
            fingerprint_score = self._calculate_fingerprint_score(doc1, doc2)
            ngram_score = self._calculate_ngram_score(doc1, doc2)
            lcs_score = self._calculate_lcs_score(doc1, doc2)
            cosine_score = self.cosine_similarity(doc1, doc2)
            
            weights = self._get_weights(ext1, ext2)
            
            final_score = (
                weights['fingerprint'] * fingerprint_score +
                weights['ngram'] * ngram_score +
                weights['lcs'] * lcs_score +
                weights['cosine'] * cosine_score
            )
            
            logger.info(f"Scores - Fingerprint: {fingerprint_score:.2f}, Ngram: {ngram_score:.2f}, "
                       f"LCS: {lcs_score:.2f}, Cosine: {cosine_score:.2f}, Final: {final_score:.2f}")
            
            return {
                'final_score': final_score,
                'fingerprint_score': fingerprint_score,
                'ngram_score': ngram_score,
                'lcs_score': lcs_score,
                'cosine_score': cosine_score,
                'is_plagiarism': final_score >= self._get_threshold(ext1, ext2)
            }
        except Exception as e:
            logger.error(f"Erreur dans calculate_similarity_score: {str(e)}")
            return {
                'final_score': 0.0,
                'fingerprint_score': 0.0,
                'ngram_score': 0.0,
                'lcs_score': 0.0,
                'cosine_score': 0.0,
                'is_plagiarism': False
            }
# Classe principale de détection de plagiat
class PlagiarismDetector:
    def __init__(self):
        self.document_processor = DocumentProcessor()
        self.similarity_engine = SimilarityEngine()
    
    def detect_plagiarism(self, file_path1: str, file_path2: str) -> Dict:
        """Détecte le plagiat entre deux fichiers"""
        logger.info(f"Analysing files: {file_path1} and {file_path2}")
        
        # Extraction des extensions
        ext1 = os.path.splitext(file_path1)[1]
        ext2 = os.path.splitext(file_path2)[1]
        
        # Traitement des documents
        try:
            doc1_text, _ = self.document_processor.process_document(file_path1)
            doc2_text, _ = self.document_processor.process_document(file_path2)
        except Exception as e:
            logger.error(f"Error processing documents: {str(e)}")
            raise
        if not doc1_text.strip():
            raise ValueError(f"Le contenu du fichier {file_path1} est vide ou non lisible.")
        if not doc2_text.strip():
            raise ValueError(f"Le contenu du fichier {file_path2} est vide ou non lisible.")
        # Calcul de la similarité
        try:
            results = self.similarity_engine.calculate_similarity_score(
                doc1_text, doc2_text, ext1, ext2
            )
        except Exception as e:
            logger.error(f"Error calculating similarity: {str(e)}")
            raise
        
        # Ajout des métadonnées
        results.update({
            'file1': file_path1,
            'file2': file_path2,
            'file1_type': ext1,
            'file2_type': ext2,
            'timestamp': '2024-01-16'
        })
        
        return results

# Générateur de rapports
class ReportGenerator:
    def generate_report(self, results: Dict) -> str:
        """Génère un rapport détaillé des résultats de la détection"""
        report = [
            "=== Rapport de Détection de Plagiat ===\n",
            f"Date d'analyse: {results['timestamp']}\n",
            f"\nFichiers analysés:",
            f"- Fichier 1: {results['file1']} ({results['file1_type']})",
            f"- Fichier 2: {results['file2']} ({results['file2_type']})\n",
            f"\nScores de similarité:",
            f"- Score Fingerprint: {results['fingerprint_score']:.2%}",
            f"- Score N-grammes: {results['ngram_score']:.2%}",
            f"- Score LCS: {results['lcs_score']:.2%}",
            f"- Score Cosinus: {results['cosine_score']:.2%}",
            f"\nScore global: {results['final_score']:.2%}",
            f"\nVerdict: {'PLAGIAT DÉTECTÉ' if results['is_plagiarism'] else 'PAS DE PLAGIAT'}"
        ]
        
        return '\n'.join(report)

# # Exemple d'utilisation
# if __name__ == "__main__":
#     # Initialisation du détecteur et du générateur de rapports
#     detector = PlagiarismDetector()
#     report_generator = ReportGenerator()
    
#     try:
#         # Détection du plagiat
#         results = detector.detect_plagiarism('document1.txt', 'document2.txt')
        
#         # Génération et affichage du rapport
#         report = report_generator.generate_report(results)
#         print(report)
        
#     except Exception as e:
#         logger.error(f"Error during plagiarism detection: {str(e)}")
#         raise